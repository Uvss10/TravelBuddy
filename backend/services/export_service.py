import os
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from datetime import datetime

def generate_itinerary_docx(itinerary_data: dict, output_path: str):
    """
    Generates a professional and attractive .docx file from itinerary JSON data.
    """
    doc = Document()

    # --- Styles ---
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Arial'
    font.size = Pt(11)

    # --- Title Section ---
    destination = itinerary_data.get('destination', 'Your Trip').upper()
    title = doc.add_heading(f"TRAVEL ITINERARY: {destination}", 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # Tagline if available
    summary = itinerary_data.get('itinerary_ai_output', {}).get('trip_summary', {})
    tagline = summary.get('tagline', 'Crafted by TravelBuddy AI')
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(tagline)
    run.italic = True
    run.font.size = Pt(12)
    run.font.color.rgb = RGBColor(0x4A, 0x90, 0xE2) # Primary blue

    doc.add_paragraph() # Spacer

    # --- Trip Metrics ---
    table = doc.add_table(rows=1, cols=3)
    table.alignment = WD_ALIGN_PARAGRAPH.CENTER
    hdr_cells = table.rows[0].cells
    
    days = itinerary_data.get('total_days', '-')
    budget = itinerary_data.get('budget_category', '-')
    style = summary.get('travel_style', 'Balanced')

    hdr_cells[0].text = f"DURATION\n{days} Days"
    hdr_cells[1].text = f"BUDGET\n{budget}"
    hdr_cells[2].text = f"STYLE\n{style}"

    for cell in hdr_cells:
        for paragraph in cell.paragraphs:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in paragraph.runs:
                run.bold = True

    doc.add_paragraph() # Spacer
    doc.add_paragraph("_" * 50).alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph() # Spacer

    # --- Day Wise Journey ---
    doc.add_heading('YOUR DAILY JOURNEY', level=1)
    
    ai_output = itinerary_data.get('itinerary_ai_output', {})
    days_list = ai_output.get('days', [])
    
    # If it's the legacy format (Day 1: [...])
    if not days_list and isinstance(ai_output, dict):
        for key, activities in ai_output.items():
            if key.startswith('Day') and isinstance(activities, list):
                doc.add_heading(key, level=2)
                for act in activities:
                    doc.add_paragraph(act, style='List Bullet')
    else:
        # Modern structured format
        for day_data in days_list:
            day_num = day_data.get('day', '?')
            doc.add_heading(f"DAY {day_num}", level=2)
            
            # Activities can be a list or specific slots
            activities = day_data.get('activities', [])
            if activities:
                for act in activities:
                    if isinstance(act, dict):
                        place = act.get('place_name', 'Activity')
                        time = act.get('recommended_time', '')
                        desc = act.get('description', '')
                        text = f"{time} - {place}" if time else place
                        p = doc.add_paragraph(style='List Bullet')
                        p.add_run(text).bold = True
                        if desc:
                            p.add_run(f"\n{desc}")
                    else:
                        doc.add_paragraph(str(act), style='List Bullet')
            else:
                # Check for morning/afternoon/evening from prompt template
                for slot in ['morning', 'afternoon', 'evening']:
                    content = day_data.get(slot)
                    if content:
                        p = doc.add_paragraph(style='List Bullet')
                        p.add_run(f"{slot.capitalize()}: ").bold = True
                        p.add_run(str(content))
                
                notes = day_data.get('notes')
                if notes:
                    p = doc.add_paragraph()
                    run = p.add_run(f"Insider Tip: {notes}")
                    run.italic = True
                    run.font.color.rgb = RGBColor(0x27, 0xAE, 0x60) # Success green

    # --- Budget Breakdown ---
    budget_data = ai_output.get('budget_breakdown', {})
    if budget_data:
        doc.add_page_break()
        doc.add_heading('ESTIMATED BUDGET', level=1)
        
        b_table = doc.add_table(rows=0, cols=2)
        items = [
            ('Accommodation', budget_data.get('accommodation')),
            ('Food & Dining', budget_data.get('food')),
            ('Local Transport', budget_data.get('transport')),
            ('Miscellaneous', budget_data.get('misc')),
        ]
        
        for item, price in items:
            row_cells = b_table.add_row().cells
            row_cells[0].text = item
            row_cells[1].text = str(price)
            row_cells[1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT

        doc.add_paragraph()
        source = budget_data.get('source', 'Verified Data')
        p = doc.add_paragraph(f"Pricing Data: {source}")
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        for run in p.runs:
            run.font.size = Pt(9)
            run.font.color.rgb = RGBColor(0x7F, 0x8C, 0x8D)

    # --- Footer ---
    section = doc.sections[0]
    footer = section.footer
    p = footer.paragraphs[0]
    p.text = f"Generated by TravelBuddy on {datetime.now().strftime('%Y-%m-%d %H:%M')} | Your companion for local adventures."
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.save(output_path)
    return output_path
